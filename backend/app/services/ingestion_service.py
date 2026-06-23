"""Ingestion orchestration: download -> extract/caption -> chunk -> contextual embed -> store.

Multimodal: images and image-bearing PDF pages are described by a vision model
(caption-then-embed), so they become retrievable through the same text pipeline.
The source image is kept in Storage and referenced from the chunk metadata.
"""
import asyncio
import base64
import io
import logging

from app.db.supabase import get_supabase_client
from app.services.chunking_service import chunk_text
from app.services.embedding_service import get_embeddings
from app.services.metadata_service import extract_metadata
from app.services.langsmith import get_traced_async_openai_client
from app.routers.settings import decrypt_value

logger = logging.getLogger(__name__)

BATCH_SIZE = 50            # embeddings per batch
MAX_CAPTION_PAGES = 20     # cap vision calls per document (cost guard)
CAPTION_TEXT_THRESHOLD = 100  # chars; a page with less text is treated as visual
CAPTION_CONCURRENCY = 4    # max concurrent vision calls
IMAGE_MIME_PREFIX = "image/"

IMAGE_DESCRIBE_SYSTEM = (
    "You are an image analyst writing dense, searchable descriptions for a document "
    "retrieval system. Describe what the image shows in specific detail. Transcribe "
    "ALL visible text verbatim. If it is a chart, graph, table, or diagram, state the "
    "data, axes, labels, and any trends. Be factual; no preamble."
)


def _get_llm_settings() -> dict | None:
    """Read the configured (vision-capable) LLM settings, or None if unconfigured."""
    supabase = get_supabase_client()
    result = supabase.table("global_settings").select(
        "llm_model, llm_base_url, llm_api_key"
    ).limit(1).maybe_single().execute()
    data = result.data if result else None
    api_key = decrypt_value(data.get("llm_api_key")) if data else None
    if not api_key:
        return None
    return {
        "model": data.get("llm_model") or "gpt-4o",
        "base_url": data.get("llm_base_url") or None,
        "api_key": api_key,
    }


async def describe_image(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """Describe an image with the configured vision model for retrieval.

    Returns dense prose + transcribed text. Empty string on failure (non-fatal).
    """
    try:
        settings = _get_llm_settings()
        if not settings:
            return ""
        client = get_traced_async_openai_client(
            base_url=settings["base_url"], api_key=settings["api_key"]
        )
        data_url = "data:%s;base64,%s" % (mime_type, base64.b64encode(image_bytes).decode())
        response = await client.chat.completions.create(
            model=settings["model"],
            messages=[
                {"role": "system", "content": IMAGE_DESCRIBE_SYSTEM},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Describe this image for document search, transcribing any text."},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                },
            ],
            max_tokens=700,
            temperature=0.0,
        )
        return (response.choices[0].message.content or "").strip()
    except Exception as e:
        logger.warning(f"Image description failed: {e}")
        return ""


async def _generate_document_context(text: str, filename: str) -> str:
    """Generate a short document-level context string (prepended to chunks before embedding)."""
    try:
        settings = _get_llm_settings()
        if not settings:
            return ""
        client = get_traced_async_openai_client(
            base_url=settings["base_url"], api_key=settings["api_key"]
        )
        preview = text[:2000]
        response = await client.chat.completions.create(
            model=settings["model"],
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Write a very brief context sentence (15-25 words) describing "
                        "what this document is about. This will be prepended to text "
                        "chunks for better search. Be factual and specific."
                    ),
                },
                {"role": "user", "content": f"Filename: {filename}\n\nContent preview:\n{preview}"},
            ],
            max_tokens=50,
            temperature=0.0,
        )
        context = (response.choices[0].message.content or "").strip().rstrip(".")
        logger.info(f"Document context for '{filename}': {context}")
        return context
    except Exception as e:
        logger.warning(f"Document context generation failed: {e}")
        return ""


def _prepend_context(chunks: list[str], context: str) -> list[str]:
    """Prepend document context to each chunk for contextual embeddings."""
    if not context:
        return chunks
    return [f"[Document context: {context}]\n\n{chunk}" for chunk in chunks]


def _upload_png(supabase, path: str, png_bytes: bytes) -> None:
    """Upload a rendered page image to the documents bucket (upsert)."""
    supabase.storage.from_("documents").upload(
        path=path,
        file=png_bytes,
        file_options={"content-type": "image/png", "upsert": "true"},
    )


def _pdf_extract_pages(file_bytes: bytes) -> list[dict]:
    """Per-page text via pypdfium2; render low-text pages to PNG for captioning.

    Returns a list of {index, text, png} where png is bytes for caption candidates
    (scanned/visual pages, capped by MAX_CAPTION_PAGES) and None otherwise.
    Blocking — run via asyncio.to_thread.
    """
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(file_bytes)
    pages_out: list[dict] = []
    caption_budget = MAX_CAPTION_PAGES
    try:
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            text = textpage.get_text_range() or ""
            png = None
            if len(text.strip()) < CAPTION_TEXT_THRESHOLD and caption_budget > 0:
                bitmap = page.render(scale=2.0)
                pil_image = bitmap.to_pil()
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG")
                png = buf.getvalue()
                caption_budget -= 1
                bitmap.close()
            pages_out.append({"index": i, "text": text, "png": png})
            textpage.close()
            page.close()
    finally:
        pdf.close()
    return pages_out


async def _segment_pdf(file_bytes: bytes, doc: dict, user_id: str) -> tuple[list[dict], bool]:
    """Split a PDF into provenance-tagged segments (text pages + captioned image pages)."""
    document_id = doc["id"]
    supabase = get_supabase_client()
    pages = await asyncio.to_thread(_pdf_extract_pages, file_bytes)

    sem = asyncio.Semaphore(CAPTION_CONCURRENCY)

    async def _caption(pd: dict) -> dict:
        async with sem:
            caption = await describe_image(pd["png"], "image/png")
        image_path = f"{user_id}/{document_id}/p{pd['index']}.png"
        try:
            await asyncio.to_thread(_upload_png, supabase, image_path, pd["png"])
        except Exception as e:
            logger.warning(f"Page image upload failed (p{pd['index']}): {e}")
            image_path = None
        return {
            "text": caption or f"[Image on page {pd['index'] + 1}]",
            "content_type": "image",
            "image_path": image_path,
            "page_number": pd["index"],
        }

    text_segments: list[dict] = []
    caption_jobs = []
    for pd in pages:
        if pd["png"] is not None:
            caption_jobs.append(_caption(pd))
        elif pd["text"].strip():
            text_segments.append({
                "text": pd["text"],
                "content_type": "text",
                "image_path": None,
                "page_number": pd["index"],
            })

    image_segments: list[dict] = []
    if caption_jobs:
        for r in await asyncio.gather(*caption_jobs, return_exceptions=True):
            if isinstance(r, dict):
                image_segments.append(r)

    segments = text_segments + image_segments
    segments.sort(key=lambda s: s["page_number"] if s["page_number"] is not None else 0)
    return segments, len(image_segments) > 0


async def _build_segments(file_bytes: bytes, file_type: str, doc: dict, user_id: str) -> tuple[list[dict], bool]:
    """Turn a file into provenance-tagged segments and whether it has image content."""
    if file_type.startswith(IMAGE_MIME_PREFIX):
        caption = await describe_image(file_bytes, file_type)
        text = caption or f"[Image: {doc['filename']}]"
        return [{
            "text": text,
            "content_type": "image",
            "image_path": doc["storage_path"],  # the uploaded image itself
            "page_number": None,
        }], True

    if file_type == "application/pdf":
        return await _segment_pdf(file_bytes, doc, user_id)

    # Text-based formats (txt/md/docx/xlsx/html/csv) — blocking extraction off-loop
    text = await asyncio.to_thread(extract_text, file_bytes, file_type)
    return [{
        "text": text,
        "content_type": "text",
        "image_path": None,
        "page_number": None,
    }], False


async def process_document(document_id: str, user_id: str) -> None:
    """Process an uploaded document: extract/caption, chunk, contextual embed, store."""
    supabase = get_supabase_client()

    try:
        supabase.table("documents").update({"status": "processing"}).eq("id", document_id).execute()

        doc_result = supabase.table("documents").select("*").eq("id", document_id).single().execute()
        doc = doc_result.data
        if not doc:
            raise ValueError(f"Document {document_id} not found")

        storage_path = doc["storage_path"]
        file_bytes = await asyncio.to_thread(
            supabase.storage.from_("documents").download, storage_path
        )

        # Build provenance-tagged segments (text + captioned images)
        segments, has_images = await _build_segments(file_bytes, doc["file_type"], doc, user_id)

        # Flatten to chunk-level units, chunking text within a segment (never across)
        units: list[tuple[str, str, str | None, int | None]] = []
        for seg in segments:
            if not seg["text"] or not seg["text"].strip():
                continue
            if seg["content_type"] == "text":
                for ch in chunk_text(seg["text"]):
                    units.append((ch, "text", seg["image_path"], seg["page_number"]))
            else:
                # Image caption: keep whole caption as a single chunk
                units.append((seg["text"], "image", seg["image_path"], seg["page_number"]))

        if not units:
            raise ValueError("No content extracted from document")

        full_text = "\n\n".join(u[0] for u in units)
        doc_context = await _generate_document_context(full_text, doc["filename"])

        total_chunks = 0
        for i in range(0, len(units), BATCH_SIZE):
            batch = units[i:i + BATCH_SIZE]
            texts_for_embedding = _prepend_context([u[0] for u in batch], doc_context)
            embeddings = await get_embeddings(texts_for_embedding, user_id=user_id)

            records = []
            for j, ((content, ctype, image_path, page_number), embedding) in enumerate(zip(batch, embeddings)):
                records.append({
                    "document_id": document_id,
                    "user_id": user_id,
                    "content": content,
                    "chunk_index": i + j,
                    "embedding": embedding,
                    "metadata": {
                        "filename": doc["filename"],
                        "chunk_index": i + j,
                        "content_type": ctype,
                        "image_path": image_path,
                        "page_number": page_number,
                        "doc_context": doc_context or None,
                    },
                })

            supabase.table("chunks").insert(records).execute()
            total_chunks += len(batch)

        supabase.table("documents").update({
            "status": "completed",
            "chunk_count": total_chunks,
            "has_images": has_images,
        }).eq("id", document_id).execute()

        logger.info(
            f"Document {document_id} processed: {total_chunks} chunks "
            f"(contextual={bool(doc_context)}, images={has_images})"
        )

        try:
            await extract_metadata(document_id, user_id)
        except Exception as meta_err:
            logger.warning(f"Metadata extraction failed for {document_id}: {meta_err}")

    except Exception as e:
        logger.error(f"Error processing document {document_id}: {e}")
        supabase.table("documents").update({
            "status": "failed",
            "error_message": str(e),
        }).eq("id", document_id).execute()


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract text from file bytes based on file type (non-image, non-PDF-image formats)."""
    if file_type in ("text/plain", "text/markdown"):
        return file_bytes.decode("utf-8")

    if file_type == "application/pdf":
        return _extract_pdf(file_bytes)

    if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(file_bytes)

    if file_type == "application/msword":
        try:
            return _extract_docx(file_bytes)
        except Exception:
            raise ValueError("Old .doc format is not supported. Please save as .docx and re-upload.")

    if file_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_excel(file_bytes)

    if file_type == "text/html":
        return _extract_html(file_bytes)

    if file_type == "text/csv":
        return _extract_csv(file_bytes)

    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    """Plain text extraction from a PDF (fallback path; main flow uses pypdfium2 segments)."""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(file_bytes)
    pages = []
    try:
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            text = textpage.get_text_range()
            if text:
                pages.append(text)
            textpage.close()
            page.close()
    finally:
        pdf.close()
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    from docx import Document

    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _extract_excel(file_bytes: bytes) -> str:
    """Extract text from Excel bytes. Converts each sheet to readable text."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"## Sheet: {sheet_name}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


def _extract_html(file_bytes: bytes) -> str:
    """Extract text from HTML bytes. Strips scripts, styles, and nav elements."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(file_bytes, "lxml")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_csv(file_bytes: bytes) -> str:
    """Extract text from CSV bytes. Returns pipe-separated rows."""
    import csv

    text = file_bytes.decode("utf-8")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(row))
    return "\n".join(rows)
